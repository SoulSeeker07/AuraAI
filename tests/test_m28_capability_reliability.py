"""
Milestone 28 End-to-End Capability Reliability Integration Tests
Location: tests/test_m28_capability_reliability.py

Tests T1–T6 tasks end-to-end through MasterOrchestrator.orchestrate():
  - T1: Plain text file creation on Desktop via FileManager
  - T2: Binary PowerPoint presentation creation via CodeAct
  - T3: Binary Word document creation via CodeAct
  - T4: Existing Word document editing via CodeAct
  - T5: Format conversion via CodeAct
  - T6: Multi-stage research -> markdown document generation DAG
"""

import os
from pathlib import Path
from dotenv import load_dotenv
import pytest

load_dotenv(dotenv_path=Path(__file__).resolve().parents[1] / ".env")

from core.orchestration.master_orchestrator import MasterOrchestrator
from desktop.native.known_folders import resolve_known_folder


@pytest.mark.asyncio
async def test_t1_plain_text_leave_letter_creation():
    """T1: Plain text creation routes to FileManager and creates real file on disk."""
    desktop_dir = Path(resolve_known_folder("desktop"))
    target_file = desktop_dir / "leave_letter.txt"

    if target_file.exists():
        target_file.unlink()

    try:
        orchestrator = MasterOrchestrator()
        goal = "Write me a leave letter and save it to my Desktop as leave_letter.txt"

        result = await orchestrator.process_request_async(goal)
        assert result.success is True

        assert target_file.exists(), f"Expected {target_file} to exist on disk"
        content = target_file.read_text(encoding="utf-8")
        assert "leave" in content.lower() or "manager" in content.lower()
        assert len(content) > 20
    finally:
        if target_file.exists():
            target_file.unlink()


@pytest.mark.asyncio
async def test_t2_pptx_presentation_creation():
    """T2: Binary presentation creation routes to CodeAct and generates valid PPTX."""
    output_path = Path("milestones_presentation.pptx").resolve()
    if output_path.exists():
        output_path.unlink()

    try:
        orchestrator = MasterOrchestrator()
        goal = "Create a PowerPoint presentation about AuraAI milestones with 3 slides as milestones_presentation.pptx"

        result = await orchestrator.process_request_async(goal)
        assert result.success is True

        assert output_path.exists(), f"Expected {output_path} to exist on disk"
        assert output_path.stat().st_size > 1024

        from pptx import Presentation

        prs = Presentation(str(output_path))
        assert len(prs.slides) >= 2
    finally:
        if output_path.exists():
            output_path.unlink()


@pytest.mark.asyncio
async def test_t3_word_document_creation():
    """T3: Binary Word document creation routes to CodeAct and generates valid DOCX."""
    output_path = Path("tasks_summary.docx").resolve()
    if output_path.exists():
        output_path.unlink()

    try:
        orchestrator = MasterOrchestrator()
        goal = "Create a Word document summarising today's tasks and save it as tasks_summary.docx"

        result = await orchestrator.process_request_async(goal)
        assert result.success is True

        assert output_path.exists(), f"Expected {output_path} to exist on disk"
        assert output_path.stat().st_size > 1024

        import docx

        doc = docx.Document(str(output_path))
        assert len(doc.paragraphs) > 0 or len(doc.tables) > 0
    finally:
        if output_path.exists():
            output_path.unlink()


@pytest.mark.asyncio
async def test_t4_word_document_edit(tmp_path):
    """T4: Word document editing opens existing file and appends new content."""
    import docx

    initial_file = tmp_path / "meeting_notes.docx"
    doc = docx.Document()
    doc.add_heading("Team Sync", level=1)
    doc.add_paragraph("Reviewed sprint backlog.")
    doc.save(str(initial_file))

    req_file = Path("meeting_notes_updated.docx").resolve()
    if req_file.exists():
        req_file.unlink()

    try:
        from codeact.executor import DynamicCodeActExecutor
        from codeact.models import CodeActRequest

        executor = DynamicCodeActExecutor()
        req = CodeActRequest(
            goal="Open meeting_notes.docx and add a new paragraph: 'Meeting rescheduled to 3pm'",
            output_filename="meeting_notes_updated.docx",
            input_files=[initial_file],
            allowed_libraries=["python-docx"],
        )
        res = executor.run(req)
        assert res.status == "success"
        assert res.output_path is not None
        assert res.output_path.exists()

        updated_doc = docx.Document(str(res.output_path))
        full_text = " ".join(p.text for p in updated_doc.paragraphs)
        assert "3pm" in full_text or "rescheduled" in full_text
    finally:
        if req_file.exists():
            req_file.unlink()


@pytest.mark.asyncio
async def test_t5_format_conversion_pdf(tmp_path):
    """T5: Format conversion from DOCX/text to PDF via pure Python fpdf2/reportlab."""
    pdf_out = Path("converted_report.pdf").resolve()
    if pdf_out.exists():
        pdf_out.unlink()

    try:
        from codeact.executor import DynamicCodeActExecutor
        from codeact.models import CodeActRequest

        executor = DynamicCodeActExecutor()
        req = CodeActRequest(
            goal="Generate a converted PDF document with title and summary",
            output_filename="converted_report.pdf",
            allowed_libraries=["fpdf2", "reportlab"],
        )
        res = executor.run(req)
        assert res.status == "success"
        assert res.output_path is not None
        assert res.output_path.exists()
        assert res.output_path.stat().st_size > 100
    finally:
        if pdf_out.exists():
            pdf_out.unlink()


@pytest.mark.asyncio
async def test_t6_research_to_file_dag():
    """T6: Multi-stage research -> document -> persist writes summary to python_news.md."""
    target_file = Path("python_news.md").resolve()
    if target_file.exists():
        target_file.unlink()

    try:
        orchestrator = MasterOrchestrator()
        goal = "Look up the latest Python release and write a one-paragraph summary to python_news.md"

        result = await orchestrator.process_request_async(goal)
        assert result.success is True

        assert target_file.exists(), f"Expected {target_file} to exist on disk"
        content = target_file.read_text(encoding="utf-8")
        assert len(content.strip()) > 20
        assert "python" in content.lower()
    finally:
        if target_file.exists():
            target_file.unlink()
