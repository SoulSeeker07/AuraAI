"""
DOCX (Word) Parser for RAG 2.0 Knowledge Intelligence

Supports:
- Section-level chunking
- Title extraction
- Paragraph extraction
"""

import logging
from typing import List, Dict, Any, Optional
from pathlib import Path

from docx import Document
from ..models import DocumentChunk, DocumentMetadata

logger = logging.getLogger(__name__)


class DOCXParser:
    """Parse DOCX files into structured chunks."""

    def __init__(self):
        self.supported_extensions = ['.docx']

    def supports(self, file_path: Path) -> bool:
        """Check if file type is supported."""
        return file_path.suffix.lower() in self.supported_extensions

    def parse(self, file_path: Path) -> List[DocumentChunk]:
        """Parse DOCX file into document chunks."""
        chunks = []
        try:
            doc = Document(file_path)

            # Extract titles and structure
            titles = []
            for para in doc.paragraphs:
                if para.style.name.startswith('Heading'):
                    titles.append(para.text)

            # Chunk by sections
            current_section = []
            current_section_title = None
            current_section_start = 0

            for i, para in enumerate(doc.paragraphs):
                # Start new section on heading
                if para.style.name.startswith('Heading'):
                    if current_section:
                        # Save previous section
                        chunk = self._create_section_chunk(
                            file_path, current_section_title, current_section,
                            current_section_start
                        )
                        chunks.append(chunk)

                    # Start new section
                    current_section = [para.text]
                    current_section_title = para.text
                    current_section_start = i
                else:
                    current_section.append(para.text)

            # Add last section
            if current_section:
                chunk = self._create_section_chunk(
                    file_path, current_section_title, current_section,
                    current_section_start
                )
                chunks.append(chunk)

            # Add table chunks if any
            for table in doc.tables:
                table_chunk = self._create_table_chunk(file_path, table)
                if table_chunk:
                    chunks.append(table_chunk)

        except Exception as e:
            logger.error(f"Error parsing DOCX file {file_path}: {e}")
            # Fall back to simple text extraction
            chunks = self._fallback_parse(doc)

        return chunks

    def extract_metadata(self, file_path: Path) -> DocumentMetadata:
        """Extract metadata from DOCX file."""
        try:
            doc = Document(file_path)

            return DocumentMetadata(
                source=str(file_path),
                file_type="docx",
                chunk_type="file",
                chunk_id=file_path.stem,
                line_start=1,
                line_end=len(doc.paragraphs) + len(doc.tables),
                paragraph_count=len(doc.paragraphs),
                table_count=len(doc.tables),
                word_count=sum(len(p.text.split()) for p in doc.paragraphs)
            )
        except Exception as e:
            logger.error(f"Error extracting metadata from DOCX file {file_path}: {e}")
            return DocumentMetadata(
                source=str(file_path),
                file_type="docx",
                chunk_type="file",
                chunk_id=file_path.stem,
                line_start=1,
                line_end=0
            )

    def _create_section_chunk(
        self,
        file_path: Path,
        title: Optional[str],
        content: List[str],
        start_line: int
    ) -> DocumentChunk:
        """Create a section chunk from title and content."""
        full_content = '\n'.join(content)
        return DocumentChunk(
            id=f"{file_path.stem}_{start_line}",
            content=full_content,
            metadata=DocumentMetadata(
                source=str(file_path),
                file_type="docx",
                chunk_type="section",
                chunk_id=title or f"section_{start_line}",
                line_start=start_line,
                line_end=start_line + len(content),
                section_title=title,
                word_count=len(full_content.split())
            ),
            embeddings=None
        )

    def _create_table_chunk(self, file_path: Path, table) -> Optional[DocumentChunk]:
        """Create a chunk from a table."""
        try:
            table_text = "\n".join([" | ".join(row.text for row in table.rows)])
            return DocumentChunk(
                id=f"{file_path.stem}_table_{table._tbl.xml[-10:]}",
                content=f"TABLE:\n{table_text}",
                metadata=DocumentMetadata(
                    source=str(file_path),
                    file_type="docx",
                    chunk_type="table",
                    chunk_id=f"table_{table._tbl.xml[-10:]}",
                    line_start=-1,
                    line_end=-1,
                    table_text=table_text,
                    rows=len(table.rows),
                    cols=len(table.columns)
                ),
                embeddings=None
            )
        except Exception as e:
            logger.error(f"Error creating table chunk: {e}")
            return None

    def _fallback_parse(self, doc) -> List[DocumentChunk]:
        """Fallback simple text extraction."""
        chunks = []
        for i, para in enumerate(doc.paragraphs):
            if para.text.strip():
                chunks.append(DocumentChunk(
                    id=f"{doc.filename}_paragraph_{i}",
                    content=para.text,
                    metadata=DocumentMetadata(
                        source=doc.filename,
                        file_type="docx",
                        chunk_type="paragraph",
                        chunk_id=f"para_{i}",
                        line_start=i,
                        line_end=i + 1,
                        paragraph_index=i,
                        word_count=len(para.text.split())
                    ),
                    embeddings=None
                ))
        return chunks
