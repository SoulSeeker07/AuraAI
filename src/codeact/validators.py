"""
Post-Execution Output File Validators
Location: src/codeact/validators.py

Validates generated artifact files outside the sandbox to confirm that
files exist, are non-empty, and conform to the expected binary/text structure.
"""

from __future__ import annotations

import csv
import json
import logging
from pathlib import Path
from typing import Any

from .models import CodeActRequest, ExecutionAttempt, ValidationResult

logger = logging.getLogger(__name__)


def _validate_pptx(path: Path) -> tuple[bool, str]:
    try:
        from pptx import Presentation

        prs = Presentation(str(path))
        slide_count = len(prs.slides)
        if slide_count == 0:
            return False, "Presentation contains 0 slides."
        return True, f"Valid presentation with {slide_count} slides."
    except Exception as exc:
        return False, f"Malformed PPTX presentation: {exc}"


def _validate_docx(path: Path) -> tuple[bool, str]:
    try:
        import docx

        doc = docx.Document(str(path))
        para_count = len(doc.paragraphs)
        table_count = len(doc.tables)
        if para_count == 0 and table_count == 0:
            return False, "Word document contains no paragraphs or tables."
        return True, f"Valid Word document with {para_count} paragraphs and {table_count} tables."
    except Exception as exc:
        return False, f"Malformed DOCX document: {exc}"


def _validate_xlsx(path: Path) -> tuple[bool, str]:
    try:
        import openpyxl

        wb = openpyxl.load_workbook(str(path), read_only=True)
        sheets = wb.sheetnames
        wb.close()
        if not sheets:
            return False, "Excel workbook contains no sheets."
        return True, f"Valid Excel workbook with sheets: {sheets}"
    except Exception as exc:
        return False, f"Malformed XLSX workbook: {exc}"


def _validate_pdf(path: Path) -> tuple[bool, str]:
    try:
        data = path.read_bytes()
        if not data.startswith(b"%PDF"):
            return False, "PDF header '%PDF' missing."
        if b"%%EOF" not in data and b"trailer" not in data:
            return False, "PDF EOF or trailer marker missing."
        return True, f"Valid PDF structure ({len(data)} bytes)."
    except Exception as exc:
        return False, f"Malformed PDF file: {exc}"


def _validate_image(path: Path) -> tuple[bool, str]:
    try:
        from PIL import Image

        with Image.open(str(path)) as img:
            img.verify()
        return True, "Valid image format verified."
    except ImportError:
        # Fallback basic header checks for PNG/JPEG
        data = path.read_bytes()
        if data.startswith(b"\x89PNG") or data.startswith(b"\xff\xd8\xff"):
            return True, "Valid PNG/JPEG header signature."
        return False, "Invalid image header signature."
    except Exception as exc:
        return False, f"Malformed image file: {exc}"


def _validate_csv(path: Path) -> tuple[bool, str]:
    try:
        with open(path, "r", encoding="utf-8", errors="replace") as f:
            reader = csv.reader(f)
            rows = list(reader)
        if not rows:
            return False, "CSV file is empty."
        return True, f"Valid CSV file with {len(rows)} rows."
    except Exception as exc:
        return False, f"Malformed CSV file: {exc}"


def _validate_json(path: Path) -> tuple[bool, str]:
    try:
        with open(path, "r", encoding="utf-8") as f:
            json.load(f)
        return True, "Valid JSON document."
    except Exception as exc:
        return False, f"Malformed JSON: {exc}"


def _validate_text(path: Path) -> tuple[bool, str]:
    try:
        content = path.read_text(encoding="utf-8")
        if not content.strip():
            return False, "Text file is blank or contains only whitespace."
        return True, f"Valid text file ({len(content)} characters)."
    except Exception as exc:
        return False, f"Could not read text file as UTF-8: {exc}"


FORMAT_VALIDATORS = {
    ".pptx": _validate_pptx,
    ".docx": _validate_docx,
    ".xlsx": _validate_xlsx,
    ".pdf": _validate_pdf,
    ".png": _validate_image,
    ".jpg": _validate_image,
    ".jpeg": _validate_image,
    ".csv": _validate_csv,
    ".json": _validate_json,
    ".md": _validate_text,
    ".txt": _validate_text,
    ".py": _validate_text,
}


def validate(
    request: CodeActRequest, staging_dir: Path, attempt: ExecutionAttempt
) -> ValidationResult:
    """
    Perform comprehensive post-execution checks on process status and generated artifacts.
    """
    checks: list[tuple[str, bool]] = []
    error_messages: list[str] = []

    # 1. Exit code check
    exit_zero = attempt.exit_code == 0
    checks.append(("exit_code_zero", exit_zero))
    if not exit_zero:
        error_messages.append(f"Script process exited with non-zero code {attempt.exit_code}.")

    # 2. Target file existence check
    output_path = staging_dir / request.output_filename
    file_exists = output_path.exists() and output_path.is_file()
    checks.append(("file_exists", file_exists))
    if not file_exists:
        error_messages.append(f"Expected output file '{request.output_filename}' was not created.")

    # 3. Non-trivial file size check
    if file_exists:
        file_size = output_path.stat().st_size
        min_size = 2 if output_path.suffix.lower() in (".txt", ".md", ".csv", ".json", ".py") else 100
        nonzero_size = file_size >= min_size
        checks.append(("nonzero_size", nonzero_size))
        if not nonzero_size:
            error_messages.append(
                f"Generated file '{request.output_filename}' is empty or too small ({file_size} bytes)."
            )

        # 4. Format structural verification
        suffix = output_path.suffix.lower()
        validator_fn = FORMAT_VALIDATORS.get(suffix, _validate_text)
        is_valid_format, format_msg = validator_fn(output_path)
        checks.append(("format_valid", is_valid_format))
        if not is_valid_format:
            error_messages.append(f"Format validation failed: {format_msg}")

    all_passed = all(ok for _, ok in checks)
    return ValidationResult(
        passed=all_passed,
        checks=checks,
        error_message=" | ".join(error_messages) if error_messages else None,
    )
