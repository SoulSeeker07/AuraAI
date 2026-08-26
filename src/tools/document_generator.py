"""
AuraAI Real Document Generator Service
======================================
Location: src/tools/document_generator.py

Creates and formats physical Microsoft Word (.docx) and Markdown (.md) documents on disk.
"""

import os
import re
import logging
from datetime import datetime
from pathlib import Path
from typing import Dict, Any, Optional

logger = logging.getLogger(__name__)

PROJECT_ROOT = Path(__file__).resolve().parents[2]
OUTPUT_DIR = PROJECT_ROOT / "Generated_Documents"


class DocumentGenerator:
    """Creates, formats, and saves physical .docx and .md document files on disk."""

    @classmethod
    def create_document(
        cls,
        title: str,
        content: str,
        filename_base: Optional[str] = None,
        author: str = "Sreekanta",
    ) -> Dict[str, Any]:
        """
        Create and save both .docx and .md files from generated content.

        Args:
            title: Document title (e.g. "Leave Application – 3 Days")
            content: Full document body text / markdown
            filename_base: Suggested filename without extension
            author: Author / User name

        Returns:
            Dict containing file paths, status, and summary
        """
        OUTPUT_DIR.mkdir(parents=True, exist_ok=True)

        if not filename_base:
            clean_title = re.sub(r"[^\w\s-]", "", title).strip()
            clean_title = re.sub(r"[-\s]+", "_", clean_title)
            filename_base = clean_title or "Generated_Document"

        docx_path = OUTPUT_DIR / f"{filename_base}.docx"
        md_path = OUTPUT_DIR / f"{filename_base}.md"

        # 1. Save Markdown version
        try:
            full_md = f"# {title}\n\n**Author:** {author}  \n**Date:** {datetime.now().strftime('%Y-%m-%d')}\n\n---\n\n{content.strip()}\n"
            md_path.write_text(full_md, encoding="utf-8")
        except Exception as e:
            logger.warning(f"[DocumentGenerator] Error saving .md file: {e}")

        # 2. Save DOCX version using python-docx
        docx_created = False
        try:
            import docx
            from docx.shared import Pt, Inches, RGBColor
            from docx.enum.text import WD_ALIGN_PARAGRAPH

            doc = docx.Document()

            # Set standard margins
            sections = doc.sections
            for section in sections:
                section.top_margin = Inches(1.0)
                section.bottom_margin = Inches(1.0)
                section.left_margin = Inches(1.0)
                section.right_margin = Inches(1.0)

            # Title
            h1 = doc.add_heading(title, level=0)
            h1.alignment = WD_ALIGN_PARAGRAPH.CENTER

            # Metadata subheader
            meta_p = doc.add_paragraph()
            meta_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            run_date = meta_p.add_run(f"Date: {datetime.now().strftime('%B %d, %Y')}\n")
            run_date.font.size = Pt(9.5)
            run_date.font.color.rgb = RGBColor(100, 116, 139)

            # Parse lines and format into docx
            lines = content.splitlines()
            in_table = False
            table_rows = []

            for line in lines:
                stripped = line.strip()
                if not stripped or stripped == "---":
                    if in_table and table_rows:
                        cls._render_docx_table(doc, table_rows)
                        table_rows = []
                        in_table = False
                    if stripped == "---":
                        doc.add_paragraph()
                    continue

                # Table detection
                if stripped.startswith("|") and stripped.endswith("|"):
                    in_table = True
                    # Skip separator rows like |---|---|
                    if re.match(r"^\|[\s\-:|]+\|$", stripped):
                        continue
                    cells = [c.strip() for c in stripped.split("|")[1:-1]]
                    table_rows.append(cells)
                    continue
                else:
                    if in_table and table_rows:
                        cls._render_docx_table(doc, table_rows)
                        table_rows = []
                        in_table = False

                # Headings
                if stripped.startswith("### "):
                    doc.add_heading(stripped[4:], level=2)
                elif stripped.startswith("## "):
                    doc.add_heading(stripped[3:], level=1)
                elif stripped.startswith("# "):
                    doc.add_heading(stripped[2:], level=0)
                elif stripped.startswith("• ") or stripped.startswith("- "):
                    p = doc.add_paragraph(style="List Bullet")
                    cls._add_formatted_runs(p, stripped[2:])
                else:
                    p = doc.add_paragraph()
                    cls._add_formatted_runs(p, stripped)

            if in_table and table_rows:
                cls._render_docx_table(doc, table_rows)

            doc.save(str(docx_path))
            docx_created = True
            logger.info(f"[DocumentGenerator] Word document saved successfully: {docx_path}")
        except Exception as e:
            logger.error(f"[DocumentGenerator] Failed to create .docx file: {e}", exc_info=True)

        return {
            "success": True,
            "title": title,
            "filename_base": filename_base,
            "docx_path": str(docx_path) if docx_created else None,
            "md_path": str(md_path),
            "directory": str(OUTPUT_DIR),
        }

    @staticmethod
    def _add_formatted_runs(paragraph, text: str):
        """Add text with bold markdown formatting to docx paragraph."""
        parts = re.split(r"(\*\*.*?\*\*)", text)
        for part in parts:
            if part.startswith("**") and part.endswith("**"):
                paragraph.add_run(part[2:-2]).bold = True
            else:
                paragraph.add_run(part)

    @staticmethod
    def _render_docx_table(doc, rows: list):
        """Render a table grid in docx."""
        if not rows:
            return
        num_cols = max(len(r) for r in rows)
        table = doc.add_table(rows=len(rows), cols=num_cols)
        table.style = "Table Grid"

        for r_idx, row in enumerate(rows):
            for c_idx, val in enumerate(row):
                if c_idx < num_cols:
                    cell = table.cell(r_idx, c_idx)
                    cell.text = val
                    if r_idx == 0:
                        for p in cell.paragraphs:
                            for r in p.runs:
                                r.bold = True
        doc.add_paragraph()
